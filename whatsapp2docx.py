import os
from docx import Document
from docx.shared import Cm
from PIL import Image as PILImage

# CONFIG
chat_file = "chat.txt"
image_width_cm = 3  # Width for inserted images

doc = Document()
doc.add_heading("WhatsApp Chat", level=1)

def insert_image(doc, img_path):
    try:
        with PILImage.open(img_path) as img:
            width_px, height_px = img.size
            dpi = 96
            scale = (image_width_cm * dpi) / width_px
            height_cm = (height_px * scale) / dpi
        doc.add_picture(img_path, width=Cm(image_width_cm), height=Cm(height_cm))
    except Exception as e:
        doc.add_paragraph(f"[Error loading image: {img_path} ({e})]")

with open(chat_file, encoding="utf-8") as f:
    for line in f:
        line = line.strip()
        if not line:
            continue

        # Check for image references like <anexo: filename.jpg>
        if "<anexo:" in line:
            start = line.find("<attachment:") + len("<attachment:")
            end = line.find(">", start)
            img_file = line[start:end].strip()
            img_path = os.path.join(os.path.dirname(chat_file), img_file)

            # Add image if file exists
            if os.path.exists(img_path):
                doc.add_paragraph(line[:line.find("<attachment:")].strip())
                insert_image(doc, img_path)
            else:
                doc.add_paragraph(f"[Missing image: {img_file}]")
        else:
            doc.add_paragraph(line)

doc.save("whatsapp_chat.docx")
print("✅ Created: whatsapp_chat.docx")
